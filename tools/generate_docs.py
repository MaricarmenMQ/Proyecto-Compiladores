from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'documentacion'
DOCS.mkdir(exist_ok=True)
PRIMARY = '0F766E'; DARK = '122033'; LIGHT = 'DCEFEB'; MUTED = '5E6D7C'; CODE = '0D1721'


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=DARK):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text)); r.bold = bold; r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, h in enumerate(headers):
        set_cell(table.rows[0].cells[idx], h, True, PRIMARY); shade(table.rows[0].cells[idx], LIGHT)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row): set_cell(cells[idx], value)
    doc.add_paragraph()
    return table

def add_code(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), CODE); pPr.append(shd)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string('B8F7DF')

def add_bullets(doc, items):
    for item in items: doc.add_paragraph(item, style='List Bullet')

def base_doc(title, subtitle):
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.72); sec.right_margin = Inches(.72)
    styles = doc.styles
    styles['Normal'].font.name = 'Arial'; styles['Normal'].font.size = Pt(10.5)
    for name, size, color in [('Title', 23, DARK), ('Heading 1', 15.5, PRIMARY), ('Heading 2', 12.5, DARK), ('Heading 3', 11, PRIMARY)]:
        st = styles[name]; st.font.name = 'Arial'; st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
    header = sec.header.paragraphs[0]; header.text = 'SAM-Lang - Lenguaje orientado a agentes'; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8); header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = sec.footer.paragraphs[0]; footer.text = 'Proyecto final de compiladores - 2026'; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.runs[0].font.size = Pt(8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor.from_string(PRIMARY)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle); r2.font.size = Pt(13); r2.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph()
    return doc

# MANUAL DE USUARIO
u = base_doc('Manual de Usuario', 'SAM-Lang Studio y compilador C++17')
u.add_heading('1. Finalidad', level=1)
u.add_paragraph('SAM-Lang permite escribir programas declarativos para crear agentes, asignarles objetivos, memoria, herramientas y permisos, definir flujos de comportamiento y establecer comunicación entre agentes. La entrega incluye una versión web autónoma y un compilador C++17.')
u.add_heading('2. Inicio rápido: versión web', level=1)
add_bullets(u, ['Abra la carpeta web y haga doble clic en index.html.', 'Seleccione un ejemplo y presione Compilar y ejecutar.', 'Revise las pestañas Tokens, AST, Runtime, TAC y SAM-VM.', 'En Probar agente creado, escriba un mensaje y presione Enviar mensaje.'])
u.add_paragraph('La versión web funciona directamente en el navegador. No requiere Node.js, conexión a internet, clave de API ni instalación de paquetes.')
u.add_heading('3. Sintaxis básica', level=1)
add_code(u, '''agente MedicoAI {
    objetivo: "Apoyar el triaje";
    inteligencia: experto;
    memoria: persistente;
    herramientas: [base_clinica];
    permisos: [leer, usar];

    flujo consulta {
        recibir paciente;
        responder "Registrar signos y sintomas";
    }
}''')
u.add_heading('4. Comunicación multiagente', level=1)
add_code(u, '''agente Analista {
    objetivo: "Analizar datos";
}

agente Medico {
    objetivo: "Tomar decisiones";
    depende_de: [Analista];
    delega: evaluar -> Analista;
}''')
u.add_paragraph('La dependencia genera una solicitud y una respuesta en el historial de comunicación. La delegación registra la tarea y el agente destino.')
u.add_heading('5. Condicionales', level=1)
add_code(u, '''flujo saludo {
    recibir mensaje;
    si mensaje == "hola" {
        responder "Hola, soy un agente SAM-Lang";
    } sino {
        responder "Solicitud recibida";
    }
}''')
u.add_heading('6. Compilación C++', level=1)
add_table(u, ['Componente', 'Requisito'], [['Sistema', 'Windows o Linux'], ['Construcción', 'CMake 3.10 o superior'], ['Compilador', 'C++17: MSVC, MinGW, GCC o Clang'], ['Web', 'Navegador moderno']])
add_code(u, 'cmake -S . -B build\ncmake --build build --config Release')
u.add_heading('7. Ejecución en consola', level=1)
add_code(u, '.\\build\\Release\\sam_compilador.exe .\\ejemplos\\comunicacion.sam --resumen --salida salida.samvm')
add_table(u, ['Opción', 'Función'], [['--sin-tokens', 'Oculta la tabla de tokens'], ['--sin-ast', 'Oculta el AST'], ['--sin-pda', 'Omite la validación PDA'], ['--traza-pda', 'Muestra la pila y transiciones'], ['--sin-runtime', 'Compila sin ejecutar agentes'], ['--sin-codigo', 'No genera TAC ni SAM-VM'], ['--salida RUTA', 'Define el archivo final'], ['--resumen', 'Oculta tokens y AST']])
u.add_heading('8. Interpretación de resultados', level=1)
add_table(u, ['Salida', 'Significado'], [['Tokens', 'Clasificación léxica con línea y columna'], ['AST', 'Estructura jerárquica del programa'], ['Semántica', 'Validación de referencias y declaraciones'], ['Runtime', 'Agentes activos, flujos y comunicación'], ['TAC', 'Código intermedio de tres direcciones'], ['SAM-VM', 'Código final de la máquina virtual académica']])
u.add_heading('9. Errores frecuentes', level=1)
add_table(u, ['Error', 'Corrección'], [['Cadena sin cerrar', 'Añadir la comilla final'], ['Falta punto y coma', 'Cerrar el campo o acción con ;'], ['Agente inexistente', 'Declarar el agente referenciado'], ['Comparación con =', 'Usar == en condicionales'], ['Runtime sin coordinador', 'Agregar coordinador: NombreAgente;']])
u.save(DOCS / 'Manual_Usuario_SAM_Lang.docx')

# MANUAL TÉCNICO
t = base_doc('Manual Técnico', 'Arquitectura e implementación de SAM-Lang')
t.add_heading('1. Arquitectura', level=1)
t.add_paragraph('El proyecto se organiza como una cadena de compilación. El front-end reconoce y valida el código; el middle-end genera y optimiza TAC; el back-end produce SAM-VM. El runtime ejecuta una semántica operacional determinista para agentes y comunicación.')
add_table(t, ['Fase', 'C++', 'Web JavaScript', 'Responsabilidad'], [
    ['Léxico', 'analisis.cpp / tokens.h', 'Lexer', 'Caracteres a tokens; comentarios y errores'],
    ['Sintaxis', 'parser.cpp / tablas_transicion.cpp', 'Parser', 'PDA, gramática y construcción del AST'],
    ['Semántica', 'semantico.cpp', 'SemanticAnalyzer', 'Símbolos, duplicados y referencias'],
    ['Runtime', 'interprete.cpp / comunicacion.cpp', 'runAgent / runtimeReport', 'Agentes, flujos, respuestas y mensajes'],
    ['TAC', 'tac.cpp', 'TACGenerator', 'Representación intermedia'],
    ['Optimización', 'optimizador.cpp', 'optimize', 'Eliminación de redundancias'],
    ['Código final', 'generador_codigo.cpp', 'toVM', 'Emisión SAM-VM']])
t.add_heading('2. Analizador léxico', level=1)
t.add_paragraph('El lexer reconoce palabras reservadas, identificadores, números enteros y decimales, cadenas, booleanos, operadores, delimitadores y comentarios. Cada token conserva línea y columna para reportes precisos.')
add_code(t, 'agente Medico { objetivo: "triaje"; }\n\nAGENTE IDENTIFICADOR LLAVE_ABRE OBJETIVO DOS_PUNTOS CADENA PUNTO_COMA LLAVE_CIERRA')
t.add_heading('3. Análisis sintáctico y PDA', level=1)
t.add_paragraph('El parser descendente recursivo implementa las producciones para programa, agente, runtime, listas, flujo, cadena de acciones, valores y condicionales. En C++, el PDA tabular actúa como verificación académica adicional y debe coincidir con el parser.')
t.add_heading('4. AST', level=1)
add_code(t, '''NodoPrograma
├── NodoAgente
│   ├── objetivo, inteligencia y memoria
│   ├── herramientas, permisos y dependencias
│   ├── NodoDelegacion[]
│   └── NodoFlujo[]
└── NodoRuntime[]''')
t.add_heading('5. Análisis semántico', level=1)
add_bullets(t, ['Registro de agentes y runtimes.', 'Detección de nombres y campos duplicados.', 'Validación de dependencias, coordinadores, coordinación, supervisión y delegaciones.', 'Detección de autorreferencias y flujos repetidos.', 'Rechazo antes del runtime cuando existen errores.'])
t.add_heading('6. Runtime y comunicación', level=1)
t.add_paragraph('El runtime activa los agentes validados, recorre sus flujos y registra una traza. Las acciones recibir y responder permiten una prueba interactiva en la versión web. Las relaciones depende_de, delega, coordina y supervisa producen mensajes en el historial.')
t.add_heading('7. Código intermedio TAC', level=1)
add_code(t, '''CREATE_AGENT Medico
SET_OBJECTIVE Medico, "Tomar decisiones"
LINK_DEPENDENCY Medico, Analista
BEGIN_FLOW Medico, diagnostico
CALL Medico, recibir -> T1
ARGUMENTS T1, paciente
END_FLOW Medico, diagnostico''')
t.add_heading('8. Optimización', level=1)
t.add_paragraph('El optimizador conserva la semántica del programa y elimina NOP, metadatos idempotentes repetidos e instrucciones consecutivas duplicadas. Se presenta un reporte comparativo del número de instrucciones antes y después.')
t.add_heading('9. Generación SAM-VM', level=1)
add_code(t, '''LOAD_AGENT Medico
SET_OBJECTIVE Medico "Tomar decisiones"
LINK Medico -> Analista
BEGIN_FLOW Medico.diagnostico
EXECUTE Medico.recibir -> T1
PUSH_ARGS T1 paciente
END_FLOW Medico.diagnostico
HALT''')
t.add_heading('10. Pruebas automáticas', level=1)
add_table(t, ['Archivo', 'Resultado'], [['medico.sam', 'Compilación, runtime y respuesta declarada'], ['comunicacion.sam', 'Dos agentes, runtime y mensajes'], ['asistente.sam', 'Evaluación condicional'], ['error_semantico.sam', 'Rechazo por dependencia inexistente']])
add_code(t, 'node tests/test_web.js\nbash tests/test_cpp.sh')
t.add_heading('11. Alcance y limitaciones', level=1)
t.add_paragraph('SAM-Lang implementa un runtime académico determinista. No se afirma que el proyecto utilice Flex o Bison: se implementan lexer y parser manuales equivalentes, mientras los archivos JFLAP documentan los autómatas. La conexión a modelos generativos es una ampliación posible, no una dependencia del producto entregado.')
t.save(DOCS / 'Manual_Tecnico_SAM_Lang.docx')

# INFORME FINAL
i = base_doc('Informe Final', 'Desarrollo de SAM-Lang como lenguaje orientado a agentes')
i.add_heading('Resumen', level=1)
i.add_paragraph('Se desarrolló SAM-Lang, un lenguaje específico de dominio para declarar agentes, definir objetivos, memoria, herramientas, permisos, restricciones y flujos, y establecer comunicación multiagente. La solución integra dos motores compatibles: un compilador C++17 y una implementación autónoma en JavaScript. Ambos realizan análisis léxico, sintáctico y semántico, construcción de AST, ejecución mediante runtime, generación de TAC, optimización y emisión de SAM-VM.')
i.add_heading('1. Planteamiento del problema', level=1)
i.add_paragraph('Los lenguajes de propósito general permiten programar agentes, pero no representan de forma directa conceptos como objetivo, memoria, herramientas, delegación, coordinación y runtime. Esto obliga a construir infraestructura adicional antes de expresar el comportamiento. SAM-Lang plantea una sintaxis declarativa centrada en esos conceptos.')
i.add_heading('2. Objetivos', level=1)
i.add_heading('2.1 Objetivo general', level=2)
i.add_paragraph('Desarrollar un lenguaje propio orientado a agentes que cubra las fases esenciales de un compilador y ofrezca una demostración ejecutable en consola y navegador.')
i.add_heading('2.2 Objetivos específicos', level=2)
add_bullets(i, ['Diseñar tokens, reglas sintácticas y estructuras del AST.', 'Implementar análisis léxico, sintáctico y semántico.', 'Construir un runtime para agentes, flujos, condicionales y mensajes.', 'Generar TAC, aplicar optimización y emitir código SAM-VM.', 'Proporcionar ejemplos, pruebas, documentación e interfaz web autónoma.'])
i.add_heading('3. Metodología', level=1)
i.add_paragraph('El desarrollo siguió una arquitectura modular. Primero se definieron tokens y autómatas. Luego se construyeron el parser y el AST. Sobre esta representación se implementaron validaciones semánticas, runtime, TAC, optimización y código final. Finalmente se reprodujo la cadena en JavaScript para cumplir la ejecución web sin depender de un servidor.')
i.add_heading('4. Diseño del lenguaje', level=1)
add_code(i, '''programa      -> (agente | runtime)+
agente        -> agente ID { campo_agente* }
runtime       -> runtime ID { campo_runtime* }
flujo         -> (flujo | flujo_dinamico) ID { paso* }
paso          -> cadena_acciones ; | condicional
condicional   -> si valor operador valor { paso* } (sino ...)?''')
i.add_heading('5. Arquitectura de la solución', level=1)
add_table(i, ['Capa', 'Producto'], [['Front-end', 'Lexer, tokens, PDA, parser, AST y semántica'], ['Ejecución', 'Runtime, evaluación de flujo y comunicación'], ['Middle-end', 'TAC y optimización'], ['Back-end', 'SAM-VM'], ['Interfaz', 'SAM-Lang Studio en navegador'], ['Calidad', 'Ejemplos y pruebas automáticas']])
i.add_heading('6. Resultados', level=1)
add_table(i, ['Criterio', 'Evidencia'], [['Código válido', 'Tres programas de ejemplo aceptados en C++ y JavaScript'], ['Errores', 'Reporte léxico, sintáctico y semántico con ubicación'], ['AST', 'Árbol C++ y visualización JSON web'], ['Comunicación', 'Mensajes por dependencia y delegación'], ['Control de flujo', 'Condicional si/sino ejecutado con mensaje del usuario'], ['Código intermedio', 'TAC con temporales, etiquetas y saltos'], ['Código final', 'Archivo .samvm y vista web'], ['Portabilidad', 'Web sin instalación y C++ multiplataforma']])
i.add_heading('7. Correspondencia con el sílabo', level=1)
add_table(i, ['Tema', 'Cobertura'], [['Unidad 1: tokens y AFD', 'Lexer C++, Lexer JS y AFD JFLAP'], ['Unidad 1: parser', 'Parser descendente y PDA tabular'], ['Unidad 2: semántica', 'Tabla de símbolos y validaciones'], ['Unidad 2: AST e intérprete', 'Nodos AST y runtime'], ['Unidad 2: control de flujo', 'Flujos y condicionales'], ['Unidad 2: TAC', 'Generador de tres direcciones'], ['Unidad 2: optimización', 'Eliminación de redundancias'], ['Unidad 2: código final', 'SAM-VM'], ['Integración', 'Consola, navegador, pruebas y documentación']])
i.add_heading('8. Conclusiones', level=1)
add_bullets(i, ['SAM-Lang integra las etapas principales de un compilador en un producto demostrable.', 'La sintaxis representa agentes y relaciones multiagente de forma explícita.', 'La versión web elimina barreras de instalación y permite observar cada etapa.', 'La coexistencia de C++ y JavaScript facilita sustentar tanto la construcción original como la migración web.'])
i.add_heading('9. Recomendaciones', level=1)
add_bullets(i, ['Añadir persistencia de memoria y un planificador concurrente.', 'Incorporar un proveedor de IA mediante configuración externa, sin exponer claves en el navegador.', 'Ampliar la gramática con variables, asignación y expresiones aritméticas.', 'Agregar más pruebas de regresión y mediciones de cobertura.'])
i.add_heading('10. Referencias', level=1)
for ref in [
    'Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. (2006). Compilers: Principles, Techniques, and Tools (2nd ed.). Pearson.',
    'Levine, J. (2009). Flex & Bison. O’Reilly Media.',
    'ISO/IEC. (2020). ISO/IEC 14882:2020 - Programming Languages - C++.',
    'Ecma International. (2025). ECMA-262: ECMAScript Language Specification.',
    'Wooldridge, M. (2009). An Introduction to MultiAgent Systems (2nd ed.). Wiley.'
]: i.add_paragraph(ref)
i.save(DOCS / 'Informe_Final_SAM_Lang.docx')
print('Documentos generados en', DOCS)
